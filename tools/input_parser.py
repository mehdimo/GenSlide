"""
Input parser node for the GenSlide LangGraph pipeline.

Responsibilities:
    - Accept raw_input as plain text, a PDF file path, or a DOCX file path
    - Extract clean, readable text from each source type
    - Normalize whitespace and remove artefacts
    - Write the result into state["parsed_text"]
    - Write a human-readable error into state["error"] on failure
      (never raise — let the graph handle routing)

Supported input_type values:
    "text"  — raw string passed directly from the Streamlit text area
    "pdf"   — filesystem path to a .pdf file uploaded via Streamlit
    "docx"  — filesystem path to a .docx file uploaded via Streamlit
"""

import re
import logging
from pathlib import Path

from pypdf import PdfReader
from docx import Document

from graph.state import AgentState

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _normalize(text: str) -> str:
    """
    Clean up extracted text:
        - Collapse runs of blank lines into a single blank line
        - Strip leading/trailing whitespace per line
        - Remove non-printable characters (except newlines and tabs)
    """
    # Remove non-printable chars except \n and \t
    text = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]", "", text)
    # Strip each line
    lines = [line.strip() for line in text.splitlines()]
    # Collapse 3+ consecutive blank lines into 2
    cleaned: list[str] = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned.append(line)
        else:
            blank_count = 0
            cleaned.append(line)
    return "\n".join(cleaned).strip()


def _parse_text(raw: str) -> str:
    return _normalize(raw)


def _parse_pdf(path: str) -> str:
    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        extracted = page.extract_text()
        if extracted:
            pages.append(extracted)
        else:
            logger.warning("PDF page %d yielded no text (possibly scanned).", i)
    if not pages:
        raise ValueError(
            "No text could be extracted from the PDF. "
            "The file may be scanned or image-based."
        )
    return _normalize("\n\n".join(pages))


def _parse_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    if not paragraphs:
        raise ValueError("No text content found in the DOCX file.")

    return _normalize("\n\n".join(paragraphs))


# ---------------------------------------------------------------------------
# LangGraph node
# ---------------------------------------------------------------------------

def parse_input_node(state: AgentState) -> AgentState:
    """
    LangGraph node: parse raw_input into clean parsed_text.

    Reads:
        state["raw_input"]   — the raw content or file path
        state["input_type"]  — "text" | "pdf" | "docx"

    Writes:
        state["parsed_text"] — cleaned extracted text
        state["error"]       — error message string on failure (else None)
    """
    raw = state.get("raw_input", "").strip()
    input_type = state.get("input_type", "text")

    if not raw:
        state["error"] = "raw_input is empty. Please provide text or upload a file."
        state["parsed_text"] = ""
        return state

    try:
        if input_type == "text":
            parsed = _parse_text(raw)

        elif input_type == "pdf":
            path = Path(raw)
            if not path.exists():
                raise FileNotFoundError(f"PDF not found at path: {raw}")
            parsed = _parse_pdf(str(path))

        elif input_type == "docx":
            path = Path(raw)
            if not path.exists():
                raise FileNotFoundError(f"DOCX not found at path: {raw}")
            parsed = _parse_docx(str(path))

        else:
            raise ValueError(
                f"Unsupported input_type '{input_type}'. "
                "Expected 'text', 'pdf', or 'docx'."
            )

        if len(parsed) < 30:
            raise ValueError(
                f"Extracted text is too short ({len(parsed)} chars). "
                "Please provide more content."
            )

        logger.info(
            "parse_input_node: parsed %d chars from input_type='%s'",
            len(parsed),
            input_type,
        )

        state["parsed_text"] = parsed
        state["error"] = None

    except Exception as exc:
        logger.exception("parse_input_node failed: %s", exc)
        state["parsed_text"] = ""
        state["error"] = f"Input parsing failed: {exc}"

    return state